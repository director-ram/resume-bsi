# filename: app.py

from flask import Flask, render_template, request, send_file
import subprocess
from docx import Document
import os

app = Flask(__name__)

# ----------------------------
# 1️⃣ Resume Sections & CAG Prompts
# ----------------------------
resume_prompts = {
    "Objective": (
        "You are an expert resume consultant with 20 years of experience. "
        "Enhance the Objective section for a professional user (working professional, student, or career switcher). "
        "Make it polished, employer-focused, and tailored to the user's field. Highlight adaptability, career goals, "
        "enthusiasm, and strengths. Choose the most relevant option if multiple paths exist. "
        "Return only the improved text starting with 'Objective:'."
    ),
    "Experience": (
        "You are an expert resume consultant with 20 years of experience. "
        "Enhance the Experience section to emphasize achievements, responsibilities, and measurable impact. "
        "Use action verbs, quantify results, and highlight transferable skills. Choose the most relevant role if multiple exist. "
        "Return only the improved text starting with 'Experience:'."
    ),
    "Skills": (
        "You are an expert resume consultant with 20 years of experience. "
        "Improve the Skills section to be concise, organized, and impressive. "
        "Group skills logically, include technical and soft skills, and remove redundancy. "
        "Return only the improved text starting with 'Skills:'."
    ),
    "Education": (
        "You are an expert resume consultant with 20 years of experience. "
        "Rewrite the Education section to clearly present degrees, certifications, and relevant coursework. "
        "Focus on what supports the user's career goals. Return only the improved text starting with 'Education:'."
    ),
    "Projects": (
        "You are an expert resume consultant with 20 years of experience. "
        "Enhance the Projects section to clearly present scope, technologies, contributions, and measurable impact. "
        "Return only the improved text starting with 'Projects:'."
    ),
    "Certifications": (
        "You are an expert resume consultant with 20 years of experience. "
        "Improve the Certifications section to highlight relevant certifications and their impact. "
        "Make it concise and professional. Return only the improved text starting with 'Certifications:'."
    ),
    "Achievements": (
        "You are an expert resume consultant with 20 years of experience. "
        "Enhance the Achievements section to highlight awards, recognitions, or accomplishments. "
        "Make it concise, quantifiable, and professional. Return only the improved text starting with 'Achievements:'."
    ),
    "Hobbies": (
        "You are an expert resume consultant with 20 years of experience. "
        "Improve the Hobbies/Interests section to be professional, relevant, and reflective of skills or traits valued in the workplace. "
        "Return only the improved text starting with 'Hobbies:'."
    ),
}

# ----------------------------
# 2️⃣ Function: Enhance Section via Gemma-3 (CAG)
# ----------------------------
def enhance_section(section_name, user_input):
    """
    CAG: Combine context prompt + user input and enhance section using Gemma3 270M via local Ollama.
    """
    if not user_input.strip():
        return ""  # skip empty input

    combined_prompt = f"{resume_prompts[section_name]}\n\nUser Input: {user_input}"
    result = subprocess.run(
        ["ollama", "run", "gemma3:270m", combined_prompt],
        capture_output=True,
        text=True
    )
    return result.stdout.strip()

# ----------------------------
# 3️⃣ Function: Save Enhanced Resume as DOCX
# ----------------------------
def save_resume_docx(enhanced_resume, filename="Enhanced_Resume.docx"):
    doc = Document()
    doc.add_heading("Enhanced Resume", 0)
    for section, text in enhanced_resume.items():
        if text:
            doc.add_heading(section, level=1)
            doc.add_paragraph(text)
    doc.save(filename)
    return filename

#