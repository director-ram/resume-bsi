from flask import Flask, request, send_file, jsonify, make_response, send_from_directory, render_template, redirect, url_for, session
from flask_cors import CORS
from groq import Groq
from docx import Document
import os
import requests  # Still needed for health check endpoint
import re
import logging
import traceback
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from io import BytesIO
import json
import asyncio

try:
    from playwright.async_api import async_playwright
    PLAYWRIGHT_AVAILABLE = True
except Exception:
    PLAYWRIGHT_AVAILABLE = False

# --------------------------------
# DOCX Generation Function
# --------------------------------
def save_resume_docx(resume_data):
    """Generate and save a DOCX file from resume data with template and color support."""
    doc = Document()
    
    # Get template and color info
    template = resume_data.get('template', 'modern')
    color = resume_data.get('color', '')
    
    # Add title with template-specific styling
    title = doc.add_heading('Resume', 0)
    
    # Add personal information
    if 'personalInfo' in resume_data and resume_data['personalInfo']:
        personal = resume_data['personalInfo']
        if personal.get('fullName'):
            doc.add_heading(personal['fullName'], level=1)
        
        # Contact information
        contact_info = []
        if personal.get('email'):
            contact_info.append(personal['email'])
        if personal.get('phone'):
            contact_info.append(personal['phone'])
        if personal.get('location'):
            contact_info.append(personal['location'])
        if personal.get('linkedin'):
            contact_info.append(personal['linkedin'])
        
        if contact_info:
            doc.add_paragraph(' • '.join(contact_info))
        
        if personal.get('summary'):
            doc.add_heading('Professional Summary', level=2)
            doc.add_paragraph(personal['summary'])
    
    # Add experience
    if 'experience' in resume_data and resume_data['experience']:
        experience_title = 'Professional Experience' if template == 'professional' else 'Work Experience'
        doc.add_heading(experience_title, level=2)
        
        for exp in resume_data['experience']:
            if exp.get('title') or exp.get('company'):
                # Job title and company
                title_text = exp.get('title', 'Position')
                company_text = exp.get('company', 'Company')
                doc.add_paragraph(f"{title_text} at {company_text}", style='Heading 3')
                
                # Dates
                if exp.get('startDate') or exp.get('endDate'):
                    start_date = exp.get('startDate', '')
                    end_date = exp.get('endDate', '')
                    if exp.get('current'):
                        date_text = f"{start_date} - Present"
                    else:
                        date_text = f"{start_date} - {end_date}" if end_date else start_date
                    doc.add_paragraph(date_text, style='Intense Quote')
                
                # Description
                if exp.get('description'):
                    doc.add_paragraph(exp['description'])
                
                doc.add_paragraph()  # Add spacing
    
    # Add education
    if 'education' in resume_data and resume_data['education']:
        education_title = 'Education' if template in ['modern', 'professional'] else 'Academic Background'
        doc.add_heading(education_title, level=2)
        
        for edu in resume_data['education']:
            if edu.get('degree') or edu.get('school'):
                # Degree and school
                degree_text = edu.get('degree', 'Degree')
                school_text = edu.get('school', 'Institution')
                field_text = edu.get('field', '')
                
                if field_text:
                    degree_line = f"{degree_text} in {field_text}"
                else:
                    degree_line = degree_text
                
                doc.add_paragraph(f"{degree_line} from {school_text}", style='Heading 3')
                
                # Graduation date and GPA
                grad_info = []
                if edu.get('graduationDate'):
                    grad_info.append(f"Graduated: {edu['graduationDate']}")
                if edu.get('gpa'):
                    grad_info.append(f"GPA: {edu['gpa']}")
                
                if grad_info:
                    doc.add_paragraph(' • '.join(grad_info), style='Intense Quote')
                
                # Additional details
                if edu.get('description'):
                    doc.add_paragraph(edu['description'])
                
                doc.add_paragraph()  # Add spacing
    
    # Add skills
    if 'skills' in resume_data and resume_data['skills']:
        skills_title = 'Core Competencies' if template == 'professional' else 'Skills'
        doc.add_heading(skills_title, level=2)
        doc.add_paragraph(resume_data['skills'])
    
    # Add projects
    if 'projects' in resume_data and resume_data['projects']:
        projects_title = 'Key Projects' if template in ['professional', 'minimal'] else 'Projects'
        doc.add_heading(projects_title, level=2)
        
        # Split projects by double newlines or bullet points
        projects_text = resume_data['projects']
        if '\n\n' in projects_text:
            projects = [p.strip() for p in projects_text.split('\n\n') if p.strip()]
        else:
            projects = [p.strip() for p in projects_text.split('\n') if p.strip()]
        
        for project in projects:
            if project:
                doc.add_paragraph(f"• {project}")
    
    # Save to temporary file with template and color info in filename
    filename = f'resume_{template}'
    if color:
        filename += f'_{color.replace("#", "")}'
    filename += '.docx'
    
    doc.save(filename)
    return filename

# --------------------------------
# Simple .env loader (no dependency)
# --------------------------------
def _load_env_file(path: str = ".env") -> None:
    try:
        if not os.path.exists(path):
            print(f"Warning: .env file not found at {path}")
            return
        
        # Try different encodings to handle BOM and encoding issues
        encodings = ['utf-8-sig', 'utf-8', 'latin-1', 'cp1252']
        content = None
        
        for encoding in encodings:
            try:
                with open(path, "r", encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            return
            
        variables_loaded = 0
        for raw in content.splitlines():
            line = raw.strip()
            if not line or line.startswith("#"):
                continue
            if "=" not in line:
                continue
            key, val = line.split("=", 1)
            key = key.strip()
            val = val.strip().strip('"').strip("'")
            if key and key not in os.environ:
                os.environ[key] = val
                variables_loaded += 1
        
        # Only print if variables were actually loaded
        if variables_loaded > 0:
            print(f"Loaded {variables_loaded} environment variables from .env")
        
    except Exception as e:
        print(f"Error loading .env file: {e}")
        # Fail silently; env can still be provided via real env vars
        pass

_load_env_file()

app = Flask(__name__)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Reduce Groq client logging verbosity
import httpx
import httpcore
httpx_logger = logging.getLogger("httpx")
httpx_logger.setLevel(logging.WARNING)
httpcore_logger = logging.getLogger("httpcore")
httpcore_logger.setLevel(logging.WARNING)
groq_logger = logging.getLogger("groq")
groq_logger.setLevel(logging.WARNING)

# Simple CORS configuration
CORS(app, resources={r"/*": {"origins": "*"}})

# Groq API Configuration
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
GROQ_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"

# Initialize Groq client
client = None
if GROQ_API_KEY:
    try:
        client = Groq(api_key=GROQ_API_KEY)
        test_response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": "Hello"}],
            max_tokens=5,
            temperature=0.7
        )
        logger.info(f" Groq API connection successful with model: {GROQ_MODEL}")
        print(f" Groq API Key loaded and tested successfully")
    except Exception as e:
        logger.error(f"Groq API connection failed: {e}")
        print(f" Groq API connection failed: {e}")
        client = None
else:
    logger.error(" No GROQ_API_KEY found")
    print("No GROQ_API_KEY found")

# Resume Enhancement Prompts
GLOBAL_RULES = [
    "Use a professional, employer-focused tone.",
    "Do not use first-person pronouns.",
    "Be concise, quantifiable, and clear.",
    "Fix grammar, avoid redundancy.",
    "Do not invent experiences or education.",
    "Do not output anything else other than the output that goes inside a particular section; no commentaries of any sort."
]

GLOBAL_RULE = "Global Resume Rules:\n" + "\n".join(f"{i + 1}. {rule}" for i, rule in enumerate(GLOBAL_RULES)) + "\n"

resume_prompts = {

    "summary": (
        "Role: Expert Resume Consultant\n"
        "Objective: Craft a highly professional, targeted Resume Summary.\n"
        "Instruction 1: Rewrite the summary in 50–70 words (2–4 concise sentences).\n"
        "Instruction 2: Start with the professional title and years of relevant experience.\n"
        "Instruction 3: Highlight top skills, key strengths, career goals, and 1–2 measurable achievements if possible.\n"
        "Instruction 4: Use strong, varied, and descriptive professional adjectives (e.g., results-driven, dynamic, dedicated, strategic, proactive, detail-oriented) to convey expertise and impact, while maintaining a factual, concise, and highly professional tone.\n"
        "Instruction 5: Tailor language to align with the targeted job description and ATS keywords.\n"
        "Instruction 6: Avoid first-person pronouns, personal opinions, or unnecessary details.\n"
        "Instruction 7: Ensure clarity, grammatical accuracy, and conciseness.\n"
        "Instruction 8: Return only the single best, polished version for the resume.\n"
        "Notes: Follow all Global Resume Rules."
    ),
    "experience": (
        "Role: Expert Resume Consultant\n"
        "Objective: Enhance the Work Experience section.\n"
        "Instruction 1: Rewrite experience in 70–120 words.\n"
        "Instruction 2: Use 3–5 bullet points per role; start each bullet with a strong action verb (e.g., Managed, Led, Improved).\n"
        "Instruction 3: Quantify achievements wherever possible; emphasize measurable outcomes and transferable skills.\n"
        "Instruction 4: Ensure ATS-friendly plain-text bullets (dash or •), avoid +, *, or markdown symbols.\n"
        "Instruction 5: Return only the single best version for the resume.\n"
        "Notes: Follow all Global Resume Rules."
    ),
    "skills": (
        "Role: Expert Resume Consultant\n"
        "Objective: Fix spelling, grammar, and capitalization in the Skills section.\n"
        "Context: The user wants their existing skills list cleaned up professionally.\n"
        "Instruction 1: Keep ALL skills exactly as provided by the user.\n"
        "Instruction 2: Fix spelling mistakes and grammar errors.\n"
        "Instruction 3: Capitalize the first letter of each skill properly.\n"
        "Instruction 4: Use proper capitalization for technology names (e.g., 'JavaScript', 'React', 'Node.js').\n"
        "Instruction 5: Do NOT add, remove, or reorder any skills.\n"
        "Instruction 6: Do NOT group, categorize, or reorganize skills.\n"
        "Instruction 7: Return as a simple comma-separated list.\n"
        "Instruction 8: Preserve the exact order and content of user's skills.\n"
        "Notes: Follow all Global Resume Rules. Only correct spelling, grammar, and capitalization."
    ),
    "education": (
        "Role: Expert Resume Consultant\n"
        "Objective: Improve the Education section.\n"
        "Instruction 1: Rewrite in 50–100 words.\n"
        "Instruction 2: Clearly present degree, university/college, dates, certifications, and relevant coursework.\n"
        "Instruction 3: Focus on what supports career goals; concise and factual sentences.\n"
        "Instruction 4: Return only the single best version for the resume.\n"
        "Notes: Follow all Global Resume Rules."
    ),
    "projects": (
        "Role: Expert Resume Consultant\n"
        "Objective: Enhance the Projects section with proper formatting and corrections.\n"
        "Context: The user wants to showcase project work professionally.\n"
        "Instruction 1: Rewrite each project in 30–60 words. Preserve multiple projects if present.\n"
        "Instruction 2: Fix ALL spelling mistakes and grammar errors.\n"
        "Instruction 3: Capitalize the first letter of each sentence and proper nouns.\n"
        "Instruction 4: Use proper punctuation and sentence structure.\n"
        "Instruction 5: Highlight scope, technologies, contributions, and results for each project.\n"
        "Instruction 6: Do NOT add new projects, tools, or metrics beyond user input.\n"
        "Instruction 7: Separate multiple projects with blank lines. Keep each project distinct.\n"
        "Instruction 8: Do NOT combine multiple projects into one.\n"
        "Instruction 9: Ensure each project description starts with a capital letter.\n"
        "Notes: Follow all Global Resume Rules. Focus on spelling, grammar, and proper capitalization."
    ),
    "certifications": (
        "Role: Expert Resume Consultant\n"
        "Objective: Refine the Certifications section.\n"
        "Context: The user wants to emphasize relevant certifications.\n"
        "Instruction 1: Rewrite in 30–60 words.\n"
        "Instruction 2: Highlight relevance and impact.\n"
        "Instruction 3: Return only the single best version for the resume.\n"
        "Notes: Follow all Global Resume Rules."
    ),
    "achievements": (
        "Role: Expert Resume Consultant\n"
        "Objective: Improve the Achievements section.\n"
        "Context: The user wants to showcase awards and recognitions.\n"
        "Instruction 1: Rewrite in 40–80 words.\n"
        "Instruction 2: Make achievements concise, quantifiable, and professional.\n"
        "Instruction 3: Return only the single best version for the resume.\n"
        "Notes: Follow all Global Resume Rules."
    ),
}
# -------------------------------
# Helper Functions
# -------------------------------
def _call_groq_chat(model_name: str, prompt: str, timeout_seconds: int = 90, stream: bool = False) -> str:
    """Call Groq Chat Completions API with a single user message and return text."""
    if not client:
        raise Exception("Groq client not initialized")
    
    try:
        completion = client.chat.completions.create(
            model=model_name,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.3,
            max_completion_tokens=1024,
            top_p=1,
            stream=stream,
            stop=None
        )
        
        if stream:
            # Handle streaming response
            content = ""
            for chunk in completion:
                if chunk.choices[0].delta.content:
                    content += chunk.choices[0].delta.content
            return content
        else:
            # Handle non-streaming response
            content = completion.choices[0].message.content or ""
            return content
        
    except Exception as e:
        logger.error(f"Groq API call failed: {e}")
        raise Exception(f"Groq API call failed: {e}")


def enhance_section(section_name, user_input):
    """Enhance a resume section using Ollama (CAG approach)"""
    if not user_input.strip():
        return user_input

    section_key = section_name.lower()
    if section_key not in resume_prompts:
        return user_input

    global_rules = (
        "Global Editing Rules:\n"
        "- Correct all spelling, grammar, and punctuation.\n"
        "- Standardize capitalization and spacing.\n"
        "- Keep original meaning; do not invent content.\n"
        "- Prefer concise, professional wording.\n"
        "- Do NOT explain what changed; output only the final improved content.\n"
    )

    combined_prompt = (
        f"{GLOBAL_RULE}\n{resume_prompts[section_key]}\n\n"
        f"User Input (verbatim; do not add new facts):\n{user_input}\n\n"
        f"Output Policy:\n"
        f"- Return ONLY the improved text for the {section_key} section.\n"
        f"- Do NOT include headings, labels, steps, notes, or explanations.\n"
        f"- Do NOT add entities, tools, dates, metrics, or claims not present above.\n"
        f"- Do NOT output placeholders; if something is missing, leave it out.\n"
        f"- Avoid repetition and filler. Keep it concise and direct.\n"
        f"- Stay within the specified word range.\n\n"
        f"Improved Content:"
    )

    try:
        if not GROQ_API_KEY:
            return user_input

        # Groq is the primary and only path now
        enhanced_text = _call_groq_chat(GROQ_MODEL, combined_prompt, timeout_seconds=90).strip()

        # Remove any common prefixes or headings
        prefixes = [
            f"{section_name.title()}:",
            f"{section_name.upper()}:",
            "Summary:",
            "Improved Content:",
            "Enhanced:",
            "Here is the improved text:",
            "Here's the enhanced version:",
            "Explanation of changes:",
            "Explanation:",
            "Notes:",
            ##"Here's a polished and concise resume summary/objective tailored for a user:"
        ]
        for p in prefixes:
            if enhanced_text.startswith(p):
                enhanced_text = enhanced_text[len(p):].strip()

        # 1) Keep only content after the last 'Improved Content:' or 'Steps' markers (case-insensitive)
        if re.search(r"improved\s*content\s*:", enhanced_text, flags=re.IGNORECASE):
            parts = re.split(r"improved\s*content\s*:", enhanced_text, flags=re.IGNORECASE)
            enhanced_text = parts[-1].strip()
        # Remove typical step-by-step blocks the model may emit
        enhanced_text = re.sub(r"(?is)^\s*(?:steps?|process|procedure)\s*:?[\s\S]*?(?=\n\s*\S|$)", "", enhanced_text).strip()

        # 2) Remove any echoed 'Global Editing Rules' block anywhere
        enhanced_text = re.sub(r"^\s*global\s+editing\s+rules:?[\s\S]*?(?:\n\s*user\s+input\s*:\s*|$)", "", enhanced_text, flags=re.IGNORECASE)

        # 3) Remove any echoed 'User Input:' header if present
        enhanced_text = re.sub(r"^\s*user\s+input\s*:\s*", "", enhanced_text, flags=re.IGNORECASE)

        # Remove any explanation/steps lines the model may add
        lines = [ln.rstrip() for ln in enhanced_text.splitlines()]
        filtered_lines = [
            ln for ln in lines
            if not ln.strip().lower().startswith((
                "explanation of changes",
                "explanation:",
                "changes:",
                "what changed:",
                "here's what i changed",
                "notes:",
                "rationale:",
                "step ",
                "step:",
                "steps:",
                "process:",
            ))
        ]
        enhanced_text = "\n".join(filtered_lines).strip()

        # Remove common placeholder patterns the model might introduce
        enhanced_text = re.sub(r"\[[^\]]*\]", "", enhanced_text)  # [Placeholder]
        enhanced_text = re.sub(r"<[^>]*>", "", enhanced_text)        # <Placeholder>
        enhanced_text = re.sub(r"\{[^}]*\}", "", enhanced_text)    # {Placeholder}

        # Preserve bullets/asterisks; trim whitespace only and hard-cap length
        enhanced_text = enhanced_text.strip()

        # Length caps per section to avoid runaway generations
        word_caps = {
            "summary": 100,
            "experience": 140,
            "education": 120,
            "projects": 200,  # increased for multiple projects
            "skills": 200,  # skills are comma-separated
        }
        cap = word_caps.get(section_key, 150)
        words = enhanced_text.split()
        if len(words) > cap:
            enhanced_text = " ".join(words[:cap])

        # Post-process projects to ensure proper capitalization
        if section_key == "projects":
            # Split into individual projects
            project_parts = enhanced_text.split('\n\n')
            capitalized_parts = []
            
            for part in project_parts:
                part = part.strip()
                if part:
                    # Capitalize first letter of the project description
                    if part and not part[0].isupper():
                        part = part[0].upper() + part[1:]
                    
                    # Ensure proper sentence structure
                    sentences = re.split(r'(?<=[.!?])\s+', part)
                    capitalized_sentences = []
                    
                    for sentence in sentences:
                        sentence = sentence.strip()
                        if sentence:
                            # Capitalize first letter of each sentence
                            if sentence and not sentence[0].isupper():
                                sentence = sentence[0].upper() + sentence[1:]
                            capitalized_sentences.append(sentence)
                    
                    part = '. '.join(capitalized_sentences)
                    # Remove double periods
                    part = re.sub(r'\.\.+', '.', part)
                    capitalized_parts.append(part)
            
            enhanced_text = '\n\n'.join(capitalized_parts)

        # Post-process summary to capitalize names properly
        if section_key == "summary":
            # Find and capitalize names (simple pattern: first letter of each word)
            words = enhanced_text.split()
            capitalized_words = []
            for word in words:
                if word and word[0].islower() and len(word) > 1:
                    # Capitalize first letter of each word
                    capitalized_words.append(word[0].upper() + word[1:])
                else:
                    capitalized_words.append(word)
            enhanced_text = ' '.join(capitalized_words)

        # Post-process experience and education: capitalize first letters of sentences
        if section_key in ("experience", "education"):
            sentences = re.split(r'(?<=[.!?])\s+', enhanced_text)
            fixed = []
            for s in sentences:
                s = s.strip()
                if not s:
                    continue
                if s and not s[0].isupper():
                    s = s[0].upper() + s[1:]
                fixed.append(s)
            enhanced_text = ' '.join(fixed)

        # Post-process experience: capitalize company names and job titles
        if section_key == "experience":
            # Split into individual experiences
            exp_parts = enhanced_text.split('\n\n')
            capitalized_parts = []
            
            for part in exp_parts:
                part = part.strip()
                if part:
                    # Look for patterns like "job title at company name"
                    # Capitalize job titles and company names
                    lines = part.split('\n')
                    capitalized_lines = []
                    
                    for line in lines:
                        line = line.strip()
                        if line:
                            # Capitalize first letter of the line
                            if line and not line[0].isupper():
                                line = line[0].upper() + line[1:]
                            
                            # Look for "at" pattern and capitalize both sides
                            if ' at ' in line.lower():
                                parts = line.split(' at ', 1)
                                if len(parts) == 2:
                                    job_title = parts[0].strip()
                                    company_name = parts[1].strip()
                                    
                                    # Capitalize job title words
                                    job_title = ' '.join(word.capitalize() for word in job_title.split())
                                    
                                    # Capitalize company name words
                                    company_name = ' '.join(word.capitalize() for word in company_name.split())
                                    
                                    line = f"{job_title} at {company_name}"
                            
                            capitalized_lines.append(line)
                    
                    part = '\n'.join(capitalized_lines)
                    capitalized_parts.append(part)
            
            enhanced_text = '\n\n'.join(capitalized_parts)

        # Post-process education: capitalize degree names and school names
        if section_key == "education":
            # Split into individual education entries
            edu_parts = enhanced_text.split('\n\n')
            capitalized_parts = []
            
            for part in edu_parts:
                part = part.strip()
                if part:
                    # Look for patterns like "degree from school name"
                    lines = part.split('\n')
                    capitalized_lines = []
                    
                    for line in lines:
                        line = line.strip()
                        if line:
                            # Capitalize first letter of the line
                            if line and not line[0].isupper():
                                line = line[0].upper() + line[1:]
                            
                            # Look for "from" pattern and capitalize both sides
                            if ' from ' in line.lower():
                                parts = line.split(' from ', 1)
                                if len(parts) == 2:
                                    degree_name = parts[0].strip()
                                    school_name = parts[1].strip()
                                    
                                    # Capitalize degree name words (handle special cases)
                                    degree_words = degree_name.split()
                                    capitalized_degree_words = []
                                    for word in degree_words:
                                        # Handle common degree abbreviations and special words
                                        if word.lower() in ['of', 'in', 'and', 'the', 'for']:
                                            capitalized_degree_words.append(word.lower())
                                        else:
                                            capitalized_degree_words.append(word.capitalize())
                                    degree_name = ' '.join(capitalized_degree_words)
                                    
                                    # Capitalize school name words
                                    school_name = ' '.join(word.capitalize() for word in school_name.split())
                                    
                                    line = f"{degree_name} from {school_name}"
                            
                            capitalized_lines.append(line)
                    
                    part = '\n'.join(capitalized_lines)
                    capitalized_parts.append(part)
            
            enhanced_text = '\n\n'.join(capitalized_parts)

        # Normalize skills output to comma-separated list without labels
        if section_key == "skills":
            # Extract all skills from user input (case-insensitive for matching)
            original_skills = [s.strip() for s in re.split(r",|\n|;", user_input) if s.strip()]
            original_skills_lower = [s.lower() for s in original_skills]
            
            # Extract skills from enhanced output
            lines = [ln.strip().lstrip("-•*") for ln in enhanced_text.splitlines() if ln.strip()]
            enhanced_skills = []
            for ln in lines:
                parts = [p.strip() for p in re.split(r",|;", ln) if p.strip()]
                for p in parts:
                    if not p.lower().startswith(("explanation", "note", "reason", "changes", "skills")):
                        enhanced_skills.append(p)
            
            # Only keep skills that match the original (case-insensitive) and preserve original order
            filtered = []
            for orig_skill in original_skills:
                # Find the enhanced version of this skill (case-insensitive match)
                found_enhanced = False
                for enhanced_skill in enhanced_skills:
                    if orig_skill.lower() == enhanced_skill.lower():
                        filtered.append(enhanced_skill)  # Use the enhanced version (fixed spelling/caps)
                        found_enhanced = True
                        break
                
                # If no enhanced version found, keep original
                if not found_enhanced:
                    filtered.append(orig_skill)
            
            enhanced_text = ", ".join(filtered)
            
            # Post-process skills to ensure proper capitalization
            skills_list = [s.strip() for s in enhanced_text.split(",") if s.strip()]
            capitalized_skills = []
            for skill in skills_list:
                # Capitalize first letter and handle common tech names
                skill = skill.strip()
                if skill:
                    # Handle common technology names
                    tech_names = {
                        'javascript': 'JavaScript',
                        'react': 'React',
                        'node.js': 'Node.js',
                        'nodejs': 'Node.js',
                        'python': 'Python',
                        'java': 'Java',
                        'c++': 'C++',
                        'c#': 'C#',
                        'html': 'HTML',
                        'css': 'CSS',
                        'sql': 'SQL',
                        'mongodb': 'MongoDB',
                        'mysql': 'MySQL',
                        'postgresql': 'PostgreSQL',
                        'aws': 'AWS',
                        'azure': 'Azure',
                        'docker': 'Docker',
                        'kubernetes': 'Kubernetes',
                        'git': 'Git',
                        'github': 'GitHub',
                        'gitlab': 'GitLab',
                        'jira': 'Jira',
                        'agile': 'Agile',
                        'scrum': 'Scrum',
                        'api': 'API',
                        'rest': 'REST',
                        'graphql': 'GraphQL',
                        'json': 'JSON',
                        'xml': 'XML',
                        'linux': 'Linux',
                        'windows': 'Windows',
                        'macos': 'macOS',
                        'ios': 'iOS',
                        'android': 'Android'
                    }
                    
                    # Check if it's a known tech name
                    if skill.lower() in tech_names:
                        skill = tech_names[skill.lower()]
                    else:
                        # Capitalize first letter
                        skill = skill[0].upper() + skill[1:] if skill else skill
                    
                    capitalized_skills.append(skill)
            
            enhanced_text = ", ".join(capitalized_skills)

        return enhanced_text if enhanced_text else user_input

    # Ollama paths disabled; only Groq errors are handled
    except requests.RequestException as e:
        print(f"Groq API error for {section_name}: {str(e)}")
        return user_input
    except Exception as e:
        logger.error("Failed to create DOCX: %s", str(e))
        traceback.print_exc()
        raise


# Routes
@app.route("/")
def index():
    # Redirect root to React app
    return redirect(url_for('serve_app_index'))

@app.route("/enhance", methods=["POST"])
def enhance_ajax():
    """Enhance a resume section."""
    if request.method == "OPTIONS":
        return "", 200

    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data received'}), 400

        section_name = data.get('section')
        content = data.get('content')
        # Optional validation data
        validation_data = data.get('validation', {})

        if not section_name or not content:
            return jsonify({'success': False, 'error': 'Missing section name or content'}), 400

        # Validate phone number if provided
        if validation_data.get('phone'):
            is_valid, error_message = _validate_phone_number(validation_data['phone'])
            if not is_valid:
                return jsonify({'success': False, 'error': error_message}), 400

        enhanced_content = enhance_section(section_name, content)

        return jsonify({'success': True, 'enhanced_content': enhanced_content, 'section': section_name})

    except Exception as e:
        return jsonify({'success': False, 'error': f'Server error: {str(e)}'}), 500

def _validate_phone_number(phone: str) -> tuple[bool, str]:
    """Validate phone number format (including country code)."""
    # Remove all non-digit characters
    digits_only = re.sub(r'\D', '', phone)
    
    if len(digits_only) < 10:
        return False, 'Phone number must have at least 10 digits (including country code)'
    elif len(digits_only) > 20:
        return False, 'Phone number cannot exceed 20 digits (including country code)'
    else:
        return True, ''

@app.route("/download")
def download():
    filename = "Enhanced_Resume.docx"
    if os.path.exists(filename):
        return send_file(filename, as_attachment=True)
    return "File not found.", 404

@app.route("/api/generate", methods=["POST"])
def api_generate_docx():
    """Generate a DOCX from provided resume data with template and color support.

    Expected JSON body:
    {
      "personalInfo": { "fullName": "...", "email": "...", "summary": "..." },
      "experience": [ { "title": "...", "company": "...", "description": "..." } ],
      "education": [ { "degree": "...", "school": "..." } ],
      "skills": "...",
      "projects": "...",
      "template": "modern|professional|minimal|elegant",
      "color": "#hexcolor",
      "enhance": true
    }
    """
    try:
        data = request.get_json(force=True, silent=True) or {}
        enhance_flag = bool(data.get("enhance", True))
        template = data.get("template", "modern")
        color = data.get("color", "")

        # Create structured resume data similar to PDF generation
        resume_data = {
            "personalInfo": data.get("personalInfo", {}),
            "experience": data.get("experience", []),
            "education": data.get("education", []),
            "skills": data.get("skills", ""),
            "projects": data.get("projects", ""),
            "template": template,
            "color": color
        }

        # Enhance content if requested
        if enhance_flag:
            if resume_data["personalInfo"].get("summary"):
                resume_data["personalInfo"]["summary"] = enhance_section("summary", resume_data["personalInfo"]["summary"])
            
            if resume_data["skills"]:
                resume_data["skills"] = enhance_section("skills", resume_data["skills"])
            
            if resume_data["projects"]:
                resume_data["projects"] = enhance_section("projects", resume_data["projects"])
            
            # Enhance experience descriptions
            for exp in resume_data["experience"]:
                if exp.get("description"):
                    exp["description"] = enhance_section("experience", exp["description"])
            
            # Enhance education descriptions
            for edu in resume_data["education"]:
                if edu.get("description"):
                    edu["description"] = enhance_section("education", edu["description"])

        filename = save_resume_docx(resume_data)
        return send_file(filename, as_attachment=True)

    except Exception as e:
        logger.error(f"DOCX generation error: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route("/health")
def health_check():
    """Check Groq API availability and model presence"""
    if not GROQ_API_KEY:
        return jsonify({'status': 'unhealthy', 'groq_available': False, 'error': 'GROQ_API_KEY not set'}), 500
    try:
        # Probe Groq API by attempting a cheap models list call
        url = "https://api.groq.com/openai/v1/models"
        headers = {"Authorization": f"Bearer {GROQ_API_KEY}"}
        r = requests.get(url, headers=headers, timeout=10)
        r.raise_for_status()
        data = r.json() if r.content else {}
        models = [m.get('id') for m in (data.get('data') or [])]
        model_available = GROQ_MODEL in (models or []) if models else True  # Some accounts may not list all
        return jsonify({'status': 'healthy', 'groq_available': True, 'model_available': model_available})
    except Exception as e:
        return jsonify({'status': 'unhealthy', 'groq_available': False, 'error': str(e)}), 500

# -------------------------------
# Plan Pages - Routes
# -------------------------------
def _get_resume_data():
    return session.setdefault("resume_data", {
        "summary": "",
        "experience": "",
        "education": "",
        "skills": "",
        "projects": "",
    })


# -------------------------------
# Multi-step Resume Flow
# -------------------------------
@app.route("/start")
def start_resume_flow():
    session["resume_data"] = _get_resume_data()
    return redirect(url_for("resume_summary"))


@app.route("/resume/summary", methods=["GET", "POST"])
def resume_summary():
    data = _get_resume_data()
    if request.method == "POST":
        data["summary"] = request.form.get("summary", "")
        session.modified = True
        return redirect(url_for("resume_experience"))
    return render_template("resume_summary.html", data=data)


@app.route("/resume/experience", methods=["GET", "POST"])
def resume_experience():
    data = _get_resume_data()
    if request.method == "POST":
        data["experience"] = request.form.get("experience", "")
        session.modified = True
        return redirect(url_for("resume_education"))
    return render_template("resume_experience.html", data=data)


@app.route("/resume/education", methods=["GET", "POST"])
def resume_education():
    data = _get_resume_data()
    if request.method == "POST":
        data["education"] = request.form.get("education", "")
        session.modified = True
        return redirect(url_for("resume_skills"))
    return render_template("resume_education.html", data=data)


@app.route("/resume/skills", methods=["GET", "POST"])
def resume_skills():
    data = _get_resume_data()
    if request.method == "POST":
        data["skills"] = request.form.get("skills", "")
        session.modified = True
        return redirect(url_for("resume_projects"))
    return render_template("resume_skills.html", data=data)


@app.route("/resume/projects", methods=["GET", "POST"])
def resume_projects():
    data = _get_resume_data()
    if request.method == "POST":
        data["projects"] = request.form.get("projects", "")
        session.modified = True
        return redirect(url_for("resume_review"))
    return render_template("resume_projects.html", data=data)


@app.route("/resume/review", methods=["GET", "POST"])
def resume_review():
    data = _get_resume_data()
    if request.method == "POST":
        enhanced = {}
        for section in ["summary", "experience", "education", "skills", "projects"]:
            enhanced[section] = enhance_section(section, data.get(section, ""))
        filename = save_resume_docx(enhanced)
        return render_template("resume_done.html", enhanced_resume=enhanced, download_file=filename)
    return render_template("resume_review.html", data=data)


@app.route("/resume/final-edit", methods=["POST"])
def resume_final_edit():
    """Handle final editing of enhanced resume content"""
    try:
        # Get the edited content from the form
        final_resume = {}
        for section in ["summary", "experience", "education", "skills", "projects"]:
            final_resume[section] = request.form.get(section, "")
        
        # Save the final edited resume
        filename = save_resume_docx(final_resume)
        
        # Return success message with download link
        return render_template("resume_done.html", 
                             enhanced_resume=final_resume, 
                             download_file=filename,
                             final_edited=True)
    except Exception as e:
        return render_template("resume_done.html", 
                             enhanced_resume={}, 
                             error=f"Error saving final resume: {str(e)}")

# -------------------------------
# Error Handlers
# -------------------------------
@app.errorhandler(404)
def not_found(error):
    return render_template('index.html'), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error'}), 500

# -------------------------------
# Frontend (Vite) - Static Serving
# -------------------------------

# Serve built React app from react-frontend/dist if present
FRONTEND_DIST = os.path.join(os.path.dirname(__file__), 'react-frontend', 'dist')

@app.route('/app')
def serve_app_index():
    if os.path.exists(os.path.join(FRONTEND_DIST, 'index.html')):
        return send_from_directory(FRONTEND_DIST, 'index.html')
    # If not built yet, fall back to legacy page
    return render_template("index.html")

@app.route('/app/<path:path>')
def serve_app_assets(path):
    asset_path = os.path.join(FRONTEND_DIST, path)
    if os.path.exists(asset_path):
        return send_from_directory(FRONTEND_DIST, path)
    # SPA fallback
    if os.path.exists(os.path.join(FRONTEND_DIST, 'index.html')):
        return send_from_directory(FRONTEND_DIST, 'index.html')
    return render_template("index.html")

# Serve built assets when index.html references /assets/* at site root
@app.route('/assets/<path:path>')
def serve_built_assets_root(path):
    assets_dir = os.path.join(FRONTEND_DIST, 'assets')
    file_path = os.path.join(assets_dir, path)
    if os.path.exists(file_path):
        return send_from_directory(assets_dir, path)
    return ("", 404)

# Serve favicons from dist if present
@app.route('/favicon.ico')
def serve_favicon_ico():
    fav = os.path.join(FRONTEND_DIST, 'favicon.ico')
    if os.path.exists(fav):
        return send_from_directory(FRONTEND_DIST, 'favicon.ico')
    return ("", 404)

@app.route('/favicon.jpg')
def serve_favicon_jpg():
    fav = os.path.join(FRONTEND_DIST, 'favicon.jpg')
    if os.path.exists(fav):
        return send_from_directory(FRONTEND_DIST, 'favicon.jpg')
    return ("", 404)

@app.route('/favicon.png')
def serve_favicon_png():
    fav = os.path.join(FRONTEND_DIST, 'favicon.png')
    if os.path.exists(fav):
        return send_from_directory(FRONTEND_DIST, 'favicon.png')
    return ("", 404)

# -------------------------------
# PDF Generation Functions
# -------------------------------
def generate_pdf(resume_data):
    """Generate a PDF resume from the provided data."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Detect template style and color
    template = (resume_data.get('template') or 'modern').lower()

    # Use custom color if provided, otherwise use template defaults
    custom_color = resume_data.get('color')
    if custom_color:
        primary_color = colors.HexColor(custom_color)
        heading_color = colors.HexColor(custom_color)
    else:
        # Use template default colors
        if template == 'modern':
            primary_color = colors.HexColor('#2563eb')
            heading_color = colors.HexColor('#1e40af')
        elif template == 'professional':
            primary_color = colors.HexColor('#0ea5e9')
            heading_color = colors.HexColor('#0369a1')
        elif template == 'minimal':
            primary_color = colors.HexColor('#374151')
            heading_color = colors.HexColor('#1f2937')
        elif template == 'elegant':
            primary_color = colors.HexColor('#7c3aed')
            heading_color = colors.HexColor('#5b21b6')
        else:  # default to modern
            primary_color = colors.HexColor('#2563eb')
            heading_color = colors.HexColor('#1e40af')

    # Create styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=24 if template == 'modern' else 22,
        spaceAfter=30,
        alignment=1,  # Center alignment
        textColor=primary_color
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16 if template in ['modern', 'professional'] else 15,
        spaceAfter=12,
        spaceBefore=20,
        textColor=heading_color,
        fontName='Helvetica-Bold',
        underlineWidth=0,
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11 if template == 'modern' else 10.5,
        spaceAfter=6,
        leading=14
    )
    
    # Build content
    story = []
    
    # Personal Information
    personal_info = resume_data.get('personalInfo', {})
    if personal_info.get('fullName'):
        story.append(Paragraph(personal_info['fullName'], title_style))
        story.append(Spacer(1, 12))
    
    # Contact Information
    contact_info = []
    if personal_info.get('email'):
        contact_info.append(personal_info['email'])
    if personal_info.get('phone'):
        contact_info.append(personal_info['phone'])
    if personal_info.get('location'):
        contact_info.append(personal_info['location'])
    if personal_info.get('linkedin'):
        contact_info.append(personal_info['linkedin'])
    if personal_info.get('github'):
        contact_info.append(personal_info['github'])
    
    if contact_info:
        story.append(Paragraph(' | '.join(contact_info), normal_style))
        story.append(Spacer(1, 20))
    
    # Professional Summary
    if personal_info.get('summary'):
        summary_title = 'PROFESSIONAL SUMMARY' if template == 'modern' else 'Professional Summary'
        story.append(Paragraph(summary_title, heading_style))
        story.append(Paragraph(personal_info['summary'], normal_style))
        story.append(Spacer(1, 12))
    
    # Work Experience
    experience = resume_data.get('experience', [])
    if experience:
        exp_title = 'WORK EXPERIENCE' if template == 'modern' else 'Professional Experience'
        story.append(Paragraph(exp_title, heading_style))
        for exp in experience:
            # Job title and company
            job_info = f"<b>{exp.get('title', '')}</b>"
            if exp.get('company'):
                job_info += f" - {exp.get('company', '')}"
            story.append(Paragraph(job_info, normal_style))
            
            # Dates
            start_date = exp.get('startDate', '')
            end_date = exp.get('endDate', '') if not exp.get('isCurrentJob', False) else 'Present'
            if start_date or end_date:
                date_str = f"{start_date} - {end_date}" if start_date and end_date else start_date or end_date
                story.append(Paragraph(f"<i>{date_str}</i>", normal_style))
            
            # Description
            if exp.get('description'):
                story.append(Paragraph(exp['description'], normal_style))
            
            story.append(Spacer(1, 12))
    
    # Education
    education = resume_data.get('education', [])
    if education:
        edu_title = 'EDUCATION' if template == 'modern' else 'Education'
        story.append(Paragraph(edu_title, heading_style))
        for edu in education:
            # Degree and school
            edu_info = f"<b>{edu.get('degree', '')}</b>"
            if edu.get('school'):
                edu_info += f" - {edu.get('school', '')}"
            story.append(Paragraph(edu_info, normal_style))
            
            # Graduation date and GPA
            grad_date = edu.get('graduationDate', '')
            gpa = edu.get('gpa', '')
            if grad_date or gpa:
                details = []
                if grad_date:
                    details.append(grad_date)
                if gpa:
                    details.append(f"GPA: {gpa}")
                story.append(Paragraph(f"<i>{' | '.join(details)}</i>", normal_style))
            
            story.append(Spacer(1, 8))
    
    # Skills
    skills = resume_data.get('skills', '')
    if skills.strip():
        skills_title = 'SKILLS' if template == 'modern' else 'Core Competencies'
        story.append(Paragraph(skills_title, heading_style))
        story.append(Paragraph(skills, normal_style))
        story.append(Spacer(1, 12))
    
    # Projects
    projects = resume_data.get('projects', '')
    if projects.strip():
        projects_title = 'PROJECTS' if template == 'modern' else 'Key Projects'
        story.append(Paragraph(projects_title, heading_style))
        # Split into separate project entries on blank lines
        project_entries = [p.strip() for p in re.split(r"\n\s*\n+", projects) if p.strip()]
        if len(project_entries) <= 1:
            # Fallback: try splitting by single newlines used as bullets
            project_entries = [p.strip() for p in re.split(r"\n[-•*]?\s*", projects) if p.strip()]
        # Build a bulleted list with proper spacing
        bullet_items = []
        for entry in project_entries:
            # Create a paragraph with proper left margin
            para_style = ParagraphStyle(
                'ProjectBullet',
                parent=normal_style,
                leftIndent=25,
                spaceAfter=8
            )
            bullet_items.append(ListItem(Paragraph(entry, para_style), leftIndent=0, bulletIndent=15))
        
        story.append(ListFlowable(bullet_items, bulletType='bullet', start='bulletchar', leftIndent=0, spaceAfter=12))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

async def _render_pdf_via_playwright(base_url: str, resume_data: dict, template_key: str) -> bytes:
    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        context = await browser.new_context()
        page = await context.new_page()
        
        # Seed localStorage
        import json
        await page.add_init_script(
            script=f"""
            (() => {{
                try {{
                    const data = {json.dumps(resume_data)};
                    localStorage.setItem('resumeData', JSON.stringify(data));
                }} catch (e) {{
                    console.error('localStorage error:', e);
                }}
            }})();
            """
        )
        
        color_param = resume_data.get('color', '')
        print(f"DEBUG: Color parameter: {color_param}")
        
        # URL encode the color parameter
        import urllib.parse
        encoded_color = urllib.parse.quote(color_param) if color_param else ''
        url = f"{base_url.rstrip('/')}/app/print?template={template_key}&color={encoded_color}"
        print(f"DEBUG: Generated URL: {url}")
        await page.goto(url, wait_until='networkidle')
        
        try:
            await page.wait_for_function("window.__previewReady === true", timeout=10000)
        except Exception:
            pass  # Continue anyway
        
        pdf = await page.pdf(format='A4', print_background=True, margin={ 'top': '12mm', 'bottom': '12mm', 'left': '12mm', 'right': '12mm' })
        await context.close()
        await browser.close()
        return pdf

@app.route("/api/generate-pdf", methods=["POST"])
def api_generate_pdf():
    """Generate a PDF from provided resume data."""
    try:
        data = request.get_json(force=True, silent=True) or {}

        # Try headless browser (pixel-perfect) first
        if PLAYWRIGHT_AVAILABLE:
            try:
                template_key = (data.get('template') or 'modern').lower()
                base_url = request.host_url
                pdf_bytes = asyncio.run(_render_pdf_via_playwright(base_url, data, template_key))
                return send_file(BytesIO(pdf_bytes), as_attachment=True, download_name=f"{data.get('personalInfo', {}).get('fullName', 'Resume')}.pdf", mimetype='application/pdf')
            except Exception:
                # Fall back to ReportLab if Playwright errors
                pass

        # Fallback: ReportLab
        pdf_buffer = generate_pdf(data)
        return send_file(pdf_buffer, as_attachment=True, download_name=f"{data.get('personalInfo', {}).get('fullName', 'Resume')}.pdf", mimetype='application/pdf')
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

# -------------------------------
# Main
# -------------------------------
if __name__ == "__main__":
    print("=" * 60)
    print("Resume Builder Server Starting")
    print("=" * 60)
    print(f"Model: {GROQ_MODEL}")
    print(f"API Key: {'Configured' if GROQ_API_KEY else 'Missing'}")
    print(f"Groq Client: {'Connected' if client else 'Failed'}")
    print(f"Server: http://localhost:5000")
    print("=" * 60)
    
    # Debug environment loading
    print(f"Environment GROQ_API_KEY loaded: {bool(GROQ_API_KEY)}")
    if GROQ_API_KEY:
        print(f"API Key length: {len(GROQ_API_KEY)}")
        print(f"API Key starts with: {GROQ_API_KEY[:10]}...")

    app.run(debug=True, host='0.0.0.0', port=5000)